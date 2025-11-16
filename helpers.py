import os
import sqlite3
from flask import redirect, session, g
from functools import wraps
from google import genai
from docx import Document

# --- Configuration ---
UPLOAD_FOLDER = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'uploads')
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt'}
DATABASE = 'database.db'

# Create upload folder if it doesn't exist
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Initialize Gemini API client
api_key = os.getenv("GEMINI_API_KEY")
try:
    client = genai.Client(api_key=api_key)
    print("Gemini client initialized successfully.")
except Exception as e:
    print(f"Error initializing Gemini client: {e}")
    client = None

# --- Database Functions ---
def init_db():
    """Initialize database with schema if it doesn't exist"""
    if os.path.exists(DATABASE):
        return
        
    print("Database not found. Creating database...")
    conn = sqlite3.connect(DATABASE)
    
    schema_file = os.path.join(os.path.dirname(__file__), 'schema.sql')
    if os.path.exists(schema_file):
        with open(schema_file, 'r') as f:
            conn.executescript(f.read())
        conn.commit()
        print("Database created successfully!")
    else:
        print(f"Warning: Schema file not found at {schema_file}")
    
    conn.close()

def get_db():
    """Get database connection"""
    db = getattr(g, '_database', None)
    if db is None:
        db = g._database = sqlite3.connect(DATABASE)
        db.row_factory = sqlite3.Row
        db.execute("PRAGMA foreign_keys = ON")
    return db

def query_db(query, args=(), one=False):
    """Execute a query and return results"""
    cur = get_db().execute(query, args)
    rv = cur.fetchall()
    cur.close()
    return (rv[0] if rv else None) if one else rv

def execute_db(query, args=()):
    """Execute a query that modifies the database"""
    db = get_db()
    cur = db.execute(query, args)
    db.commit()
    lastrowid = cur.lastrowid
    cur.close()
    return lastrowid

def add_syllabus_result(user_id, name, summary, resources, semester_start_date=None, semester_end_date=None):
    """
    Store processed syllabus data in the results table.
    
    Args:
        user_id (int): The ID of the user
        name (str): Name/title of the course
        summary (str): AI-generated summary of the syllabus
        resources (str): AI-generated resources in markdown format
        semester_start_date (str, optional): Start date of semester (YYYY-MM-DD)
        semester_end_date (str, optional): End date of semester (YYYY-MM-DD)
    
    Returns:
        int: The ID of the inserted record
    """
    try:
        from datetime import datetime
        current_date = datetime.now().strftime('%Y-%m-%d')
        
        result_id = execute_db(
            '''INSERT INTO results 
               (user_id, name, summary, resources, semester_start_date, semester_end_date, current_date) 
               VALUES (?, ?, ?, ?, ?, ?, ?)''',
            [user_id, name, summary, resources, semester_start_date, semester_end_date, current_date]
        )
        print(f"Syllabus result added to database with ID: {result_id}")
        return result_id
    except Exception as e:
        print(f"Error adding syllabus result to database: {e}")
        return None

def get_user_results(user_id):
    """
    Retrieve all processed syllabus results for a user from the database.
    
    Args:
        user_id (int): The ID of the user
    Returns:
        list: List of result entries for the user, or empty list if none exist
    """
    try:
        rows = query_db('SELECT * FROM results WHERE user_id = ? ORDER BY current_date DESC', [user_id])
        return [dict(row) for row in rows]
    except Exception as e:
        print(f"Error retrieving results: {e}")
        return []
    
# --- Decorators ---
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if session.get("user_id") is None:
            return redirect("/login")
        return f(*args, **kwargs)
    return decorated_function

# --- Helper Functions ---
def allowed_file(filename):
    """Check if the file extension is allowed"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def ai_analyze_file(filepath):
    """Upload and analyze file with Gemini API"""
    if not client:
        return "API client not initialized. Cannot proceed."

    uploaded_file = None
    file_ext = filepath.lower().rsplit('.', 1)[-1]
    config = {"temperature": 0.0}

    try:
        # Handle DOCX files by extracting text
        if file_ext in ['docx', 'doc']:
            print(f"\n[1/3] Extracting text from DOCX: {os.path.basename(filepath)}...")
            doc = Document(filepath)
            text_content = '\n'.join([p.text for p in doc.paragraphs])
            print("[2/3] Text extracted successfully.")
            
            print("[3/3] Requesting analysis from gemini-2.5-flash...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=["Give me a concise summary of this syllabus (start immediately with the summary, no preamble):\n\n" + text_content],
                config=config
            )
            summary = response.text
            print("[3/3] Summary received.")
        
        # Handle PDF and TXT files by uploading
        else:
            print(f"\n[1/4] Uploading file: {os.path.basename(filepath)}...")
            uploaded_file = client.files.upload(file=filepath)
            print(f"[2/4] File uploaded. URI: {uploaded_file.uri}")

            print("[3/4] Requesting analysis from gemini-2.5-flash...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=["Give me a concise summary of this syllabus (start immediately with the summary, no preamble).", uploaded_file],
                config=config
            )
            
            summary = response.text
            print("[4/4] Summary received.")
        
        return summary

    except FileNotFoundError:
        return f"Error: File not found at path: {filepath}"
    except Exception as e:
        return f"An unexpected error occurred: {e}"
    finally:
        # Clean up uploaded file resource
        if uploaded_file:
            print(f"\n[CLEANUP] Deleting uploaded file: {uploaded_file.name}")
            try:
                client.files.delete(name=uploaded_file.name)
                print("[CLEANUP] File deleted.")
            except Exception as e:
                print(f"[CLEANUP] Failed to delete: {e}")

def ai_validate_syllabus(filepath):
    """Check if uploaded file is a syllabus using Gemini API"""
    if not client:
        return False, "API client not initialized."

    uploaded_file = None
    file_ext = filepath.lower().rsplit('.', 1)[-1]

    try:
        # Handle DOCX files by extracting text
        if file_ext in ['docx', 'doc']:
            print(f"\n[VALIDATION] Extracting text from DOCX: {os.path.basename(filepath)}...")
            doc = Document(filepath)
            text_content = '\n'.join([p.text for p in doc.paragraphs])[:2000]  # Only first 2000 chars
            
            print("[VALIDATION] Checking if document is a syllabus...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=[f"Is this a course syllabus or curriculum document? Answer only 'yes' or 'no':\n\n{text_content}"]
            )
            answer = response.text.strip().lower()
        
        # Handle PDF and TXT files by uploading
        else:
            print(f"\n[VALIDATION] Uploading file: {os.path.basename(filepath)}...")
            uploaded_file = client.files.upload(file=filepath)
            
            print("[VALIDATION] Checking if document is a syllabus...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=["Is this a course syllabus or curriculum document? Answer only 'yes' or 'no'.", uploaded_file]
            )
            answer = response.text.strip().lower()
        
        is_valid = 'yes' in answer
        print(f"[VALIDATION] Result: {'Valid syllabus' if is_valid else 'Not a syllabus'}")
        return is_valid, "Valid syllabus" if is_valid else "This does not appear to be a syllabus"

    except Exception as e:
        print(f"[VALIDATION] Error: {e}")
        return False, f"Validation error: {e}"
    finally:
        if uploaded_file:
            try:
                client.files.delete(name=uploaded_file.name)
            except:
                pass

def ai_generate_resources(filepath):
    """Generate learning resources based on syllabus content"""
    if not client:
        print("[ERROR] API client not initialized. Cannot proceed.")
        return "API client not initialized. Cannot proceed."
    uploaded_file = None
    file_ext = filepath.lower().rsplit('.', 1)[-1]
    config = {"temperature": 0.0}

    try:
        print(f"\n[START] Generating resources for: {os.path.basename(filepath)}")

        # Prepare prompt template with explicit markdown formatting request
        prompt_intro = """Using the information in the syllabus I will send,
            generate a markdown-formatted list of:
            - Course resources (with authors)
            - Course instructors
            - Recommended supplementary resources (books, articles, videos, websites with URLs)
            (i.e.: youtube playlists, online courses, etc. also find some not mentioned in the syllabus)
            
            Format your response as a markdown list using bullet points (-).
            Include links in markdown format: [Title](URL) when URLs are available.
            (start immediately with the markdown list, no preamble)
        """

        # Handle DOCX files by extracting text
        if file_ext in ['docx', 'doc']:
            print("[1/3] Extracting text from DOCX...")
            doc = Document(filepath)
            text_content = '\n'.join([p.text for p in doc.paragraphs])
            print("[2/3] Text extracted successfully. Preparing prompt...")
            text_for_prompt = text_content[:15000]
            contents = [prompt_intro + "\n\nSyllabus content:\n" + text_for_prompt]

            print("[3/3] Requesting resources from gemini-2.5-flash...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=contents,
                config=config
            )
            result_text = response.text
            print("[DONE] Resources received for DOCX.")

        # Handle PDF and TXT files by uploading
        else:
            print("[1/4] Uploading file...")
            uploaded_file = client.files.upload(file=filepath)
            print(f"[2/4] File uploaded. URI: {getattr(uploaded_file, 'uri', 'unknown')}")

            contents = [
                prompt_intro + "\n\nRefer to the uploaded file for the syllabus content.",
                uploaded_file
            ]

            print("[3/4] Requesting resources from gemini-2.5-flash...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=contents,
                config=config
            )
            result_text = response.text
            print("[4/4] Resources received for uploaded file.")

        return result_text.strip()

    except FileNotFoundError:
        print(f"[ERROR] File not found at path: {filepath}")
        return f"Error: File not found at path: {filepath}"
    except Exception as e:
        print(f"[ERROR] An unexpected error occurred: {e}")
        return f"An unexpected error occurred: {e}"
    finally:
        if uploaded_file:
            print(f"\n[CLEANUP] Deleting uploaded file: {getattr(uploaded_file, 'name', 'unknown')}")
            try:
                client.files.delete(name=uploaded_file.name)
                print("[CLEANUP] File deleted.")
            except Exception as e:
                print(f"[CLEANUP] Failed to delete: {e}")

def ai_generate_ics(filepath, course_name, semester_start_date=None, semester_end_date=None):
    """Generate ICS calendar file based on syllabus content"""
    if not client:
        print("[ERROR] API client not initialized. Cannot proceed.")
        return None
    
    uploaded_file = None
    file_ext = filepath.lower().rsplit('.', 1)[-1]
    config = {"temperature": 0.0}

    try:
        print(f"\n[START] Generating ICS calendar for: {os.path.basename(filepath)}")

        # Prepare prompt for extracting schedule information
        prompt_intro = f"""Analyze this syllabus and extract the weekly schedule/course timeline.
        
Course Name: {course_name}
Semester Start: {semester_start_date or 'Not specified'}
Semester End: {semester_end_date or 'Not specified'}

Generate an ICS (iCalendar) format file with:
- Weekly class meetings (if schedule is mentioned)
- Important deadlines (assignments, exams, projects)
- Office hours (if mentioned)

Format your response ONLY as valid ICS file content. Start with BEGIN:VCALENDAR and end with END:VCALENDAR.
Use VEVENT for each event. Include DTSTART, DTEND, SUMMARY, DESCRIPTION, and LOCATION when available.
For recurring weekly classes, use RRULE with FREQ=WEEKLY.
Use proper ICS date format (YYYYMMDDTHHMMSS).

Start immediately with the ICS content, no preamble or explanation.
"""

        # Handle DOCX files by extracting text
        if file_ext in ['docx', 'doc']:
            print("[1/3] Extracting text from DOCX...")
            doc = Document(filepath)
            text_content = '\n'.join([p.text for p in doc.paragraphs])
            print("[2/3] Text extracted successfully. Preparing prompt...")
            text_for_prompt = text_content[:15000]
            contents = [prompt_intro + "\n\nSyllabus content:\n" + text_for_prompt]

            print("[3/3] Requesting ICS generation from gemini-2.5-flash...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=contents,
                config=config
            )
            ics_content = response.text
            print("[DONE] ICS content received for DOCX.")

        # Handle PDF and TXT files by uploading
        else:
            print("[1/4] Uploading file...")
            uploaded_file = client.files.upload(file=filepath)
            print(f"[2/4] File uploaded. URI: {getattr(uploaded_file, 'uri', 'unknown')}")

            contents = [
                prompt_intro + "\n\nRefer to the uploaded file for the syllabus content.",
                uploaded_file
            ]

            print("[3/4] Requesting ICS generation from gemini-2.5-flash...")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=contents,
                config=config
            )
            ics_content = response.text
            print("[4/4] ICS content received for uploaded file.")

        # Clean up markdown code blocks if present
        ics_content = ics_content.strip()
        if ics_content.startswith('```'):
            lines = ics_content.split('\n')
            ics_content = '\n'.join(lines[1:-1]) if len(lines) > 2 else ics_content
        
        # Validate basic ICS structure
        if 'BEGIN:VCALENDAR' in ics_content and 'END:VCALENDAR' in ics_content:
            print("[VALIDATION] ICS content appears valid.")
            return ics_content.encode('utf-8')
        else:
            print("[WARNING] Generated content may not be valid ICS format.")
            return ics_content.encode('utf-8')

    except Exception as e:
        print(f"[ERROR] An unexpected error occurred during ICS generation: {e}")
        return None
    finally:
        if uploaded_file:
            print(f"\n[CLEANUP] Deleting uploaded file: {getattr(uploaded_file, 'name', 'unknown')}")
            try:
                client.files.delete(name=uploaded_file.name)
                print("[CLEANUP] File deleted.")
            except Exception as e:
                print(f"[CLEANUP] Failed to delete: {e}")
